"""
doc_generator.py — Luminary Enterprise Corporate DOCX Generator
================================================================
Creates publication-quality, branded DOCX deliverables with corporate
color palettes (navy blue headings, white/charcoal body), structured
Markdown parsing, tables, headers/footers, and clean layout geometry.
"""

import os
import re
from pathlib import Path
from typing import Optional, List, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn


class DocGenerator:
    """
    Corporate DOCX generator converting structured markdown or plain text
    into professional investor-ready Word documents.
    """
    def __init__(self, primary_color: str = "#003366", secondary_color: str = "#FF5500"):
        self.primary_color = self._hex_to_rgb(primary_color)
        self.secondary_color = self._hex_to_rgb(secondary_color)
        self.body_color = RGBColor(34, 34, 34)       # Charcoal
        self.muted_color = RGBColor(100, 116, 139)   # Slate gray
        self.light_bg = "F3F4F6"                     # Light gray for table alt rows
        self.navy_hex = "003366"                     # Navy hex for table headers

    @staticmethod
    def _hex_to_rgb(hex_str: str) -> RGBColor:
        h = hex_str.lstrip("#")
        if len(h) != 6:
            return RGBColor(0, 51, 102)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def _set_cell_bg(self, cell, fill_hex: str):
        """Sets cell background color via XML shading."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def _set_cell_margins(self, cell, top: int = 120, bottom: int = 120, left: int = 180, right: int = 180):
        """Sets cell internal padding in dxa (1 pt = 20 dxa)."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'<w:top w:w="{top}" w:type="dxa"/>'
            f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'<w:left w:w="{left}" w:type="dxa"/>'
            f'<w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    def _add_formatted_text(self, paragraph, text: str, font_name: str = "Calibri", font_size: int = 11, color: Optional[RGBColor] = None, default_bold: bool = False):
        """Parses inline **bold** and *italic* markdown syntax within text."""
        color = color or self.body_color
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**') and len(token) >= 4:
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(token)
                run.bold = default_bold

            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = color

    def _render_table(self, doc: Document, rows_data: List[List[str]]):
        """Renders styled corporate table with navy header and striped rows."""
        if not rows_data or not rows_data[0]:
            return

        num_cols = len(rows_data[0])
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Header row
        header_row = table.rows[0]
        trPr = header_row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for c_idx, cell_text in enumerate(rows_data[0]):
            cell = header_row.cells[c_idx]
            self._set_cell_bg(cell, self.navy_hex)
            self._set_cell_margins(cell, top=160, bottom=160, left=180, right=180)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(cell_text.strip())
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Body rows
        for r_idx, row_values in enumerate(rows_data[1:], start=1):
            row = table.rows[r_idx]
            bg_hex = self.light_bg if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_values):
                if c_idx >= num_cols:
                    break
                cell = row.cells[c_idx]
                self._set_cell_bg(cell, bg_hex)
                self._set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                self._add_formatted_text(p, val.strip(), font_name="Calibri", font_size=10)

        # Spacing after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    def generate(self, content_text: str, output_path: str, title: Optional[str] = None) -> str:
        """
        Parses structured Markdown and compiles a corporate DOCX document.
        """
        doc = Document()

        # Page Geometry & Margins (1.0 inch)
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

        # Base Styles Configuration
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = self.body_color
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = Pt(16)

        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("LUMINARY AI ENTERPRISE DELIVERABLE | CONFIDENTIAL")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = self.muted_color

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Prepared with Luminary Autonomous Agency Infrastructure • Page 1")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = self.muted_color

        lines = content_text.split('\n')
        in_table = False
        table_rows = []

        for line in lines:
            line_stripped = line.strip()

            # Handle Table
            if line_stripped.startswith('|'):
                if '---' in line_stripped:
                    continue
                cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
                if any(cells):
                    table_rows.append(cells)
                    in_table = True
                continue
            elif in_table and table_rows:
                self._render_table(doc, table_rows)
                table_rows = []
                in_table = False

            if not line_stripped:
                continue

            # Strip internal debug metadata
            if re.match(r'^(Primary|Secondary|Accent|Background|Text)\s*Color:', line_stripped, re.I):
                continue
            if line_stripped.startswith('<clarify>') or line_stripped.startswith('<suggest>'):
                continue

            # Heading 1: # Title
            if line_stripped.startswith('# '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                self._add_formatted_text(p, line_stripped[2:].strip(), font_name="Calibri", font_size=18, color=self.primary_color, default_bold=True)

            # Heading 2: ## Section
            elif line_stripped.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                self._add_formatted_text(p, line_stripped[3:].strip(), font_name="Calibri", font_size=14, color=self.primary_color, default_bold=True)

            # Heading 3: ### Subsection
            elif line_stripped.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                self._add_formatted_text(p, line_stripped[4:].strip(), font_name="Calibri", font_size=12, color=self.secondary_color, default_bold=True)

            # Bullet List Items
            elif line_stripped.startswith(('- ', '* ', '• ')):
                bullet_content = line_stripped[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                self._add_formatted_text(p, bullet_content, font_name="Calibri", font_size=11)

            # Numbered List Items
            elif re.match(r'^\d+\.\s', line_stripped):
                num_content = re.sub(r'^\d+\.\s*', '', line_stripped).strip()
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(4)
                self._add_formatted_text(p, num_content, font_name="Calibri", font_size=11)

            # Regular Paragraph
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                self._add_formatted_text(p, line_stripped, font_name="Calibri", font_size=11)

        if in_table and table_rows:
            self._render_table(doc, table_rows)

        out_file = Path(output_path)
        out_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_file))
        return str(out_file)


doc_generator = DocGenerator()


def generate_doc_file(content_text: str, output_path: str, prompt: str = "") -> str:
    return doc_generator.generate(content_text, output_path)
